import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from scipy import stats

# Data
data = [13, 15, 16, 16, 19, 20, 20, 21, 22, 22, 25, 25, 25, 25, 30, 33, 33, 35, 35, 35, 35, 36, 40, 45, 46, 52, 70]

# (a) Mean and Median
mean = np.mean(data)
median = np.median(data)

# (b) Mode and Modality
mode_result = stats.mode(data)
mode = mode_result.mode
mode_count = mode_result.count

if mode_count == 1:
    modality = "unimodal"
elif mode_count == 2:
    modality = "bimodal"
else:
    modality = "no mode"

# (c) Midrange
midrange = (max(data) + min(data)) / 2

# (d) Roughly estimate Q1 and Q3
Q1 = np.percentile(data, 25)
Q3 = np.percentile(data, 75)

# (e) Five-number summary
minimum = min(data)
maximum = max(data)

# (f) Boxplot
plt.boxplot(data)
plt.title("Boxplot of Age Data")
plt.ylabel("Age")
plt.savefig("boxplot.png")
plt.close()

# Create a DOCX document
doc = Document()

# Add answers to the document
doc.add_heading("Data Analysis", 0)

doc.add_heading("(a) Mean and Median", level=1)
doc.add_paragraph(f"Mean: {mean}")
doc.add_paragraph(f"Median: {median}")

doc.add_heading("(b) Mode and Modality", level=1)
doc.add_paragraph(f"Mode: {mode}, Modality: {modality}")

doc.add_heading("(c) Midrange", level=1)
doc.add_paragraph(f"Midrange: {midrange}")

doc.add_heading("(d) Roughly Estimated Quartiles (Q1 and Q3)", level=1)
doc.add_paragraph(f"Q1 (Approximate): {Q1}")
doc.add_paragraph(f"Q3 (Approximate): {Q3}")

doc.add_heading("(e) Five-number Summary", level=1)
doc.add_paragraph(f"Minimum: {minimum}")
doc.add_paragraph(f"Q1 (Approximate): {Q1}")
doc.add_paragraph(f"Median: {median}")
doc.add_paragraph(f"Q3 (Approximate): {Q3}")
doc.add_paragraph(f"Maximum: {maximum}")

# Add the boxplot to the document
doc.add_heading("(f) Boxplot", level=1)
doc.add_picture("boxplot.png", width=Inches(5), height=Inches(3))

# Save the document
doc.save("data_analysis.docx")

print("Data analysis and boxplot saved to 'data_analysis.docx'.")
